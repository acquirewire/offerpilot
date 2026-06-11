"""Document tailoring (Module 2): CV bullets + cover letters, layout-preserving.

Strategy for "same layout, new words":
  * Load the master .docx with python-docx and rewrite text *in place*, one
    paragraph at a time, keeping each paragraph's style and the first run's font.
    We never rebuild the document, so headers, columns, spacing and fonts are
    exactly the original's.
  * Every rewritten bullet is forced through the XYZ formula
    ("Accomplished [X] as measured by [Y], by doing [Z]") by the prompt.
  * Cover letters get tone-matched to the JD/company voice.

Heavy deps (anthropic, python-docx) are imported lazily so the rest of the suite
runs without them. PDF export shells out to LibreOffice (`soffice`) when present.
"""
from __future__ import annotations

import json
import os
import re
import shutil
import subprocess
from dataclasses import dataclass

MODEL = "claude-opus-4-8"

_XYZ_SYSTEM = """You rewrite finance-CV bullet points to win a SPECIFIC role.
Rules, non-negotiable:
- Re-ANGLE the bullet to foreground the aspect most relevant to the target job —
  change the emphasis, structure and framing, not just the wording. (e.g. for a
  debt/restructuring role, lead with the credit, leverage or capital-structure
  side of the work; for a research role, lead with the analysis/thesis.) Stay
  strictly truthful to what the candidate actually did — never invent experience.
- Follow the XYZ idea (what you did, measured by a result, by doing how) as far as
  the length allows; lead with a strong past-tense verb.
- Preserve the key real number/metric; never invent figures.
- Mirror the job description's terminology where truthful.
Return ONLY the rewritten bullet text, nothing else."""

_COVER_SYSTEM = """You write concise, specific cover-letter paragraphs for finance roles.
Match the company's tone (inferred from the job description): if the JD is formal
and understated, be formal and understated; if energetic, be warmer. Never use
clichés ("I am writing to express my interest"), never invent experience, and
keep to the candidate's real background. Return ONLY the paragraph text."""


def _client():
    """Lazy Anthropic client with prompt caching enabled by default."""
    from anthropic import Anthropic

    return Anthropic(api_key=os.environ["ANTHROPIC_API_KEY"])


def _complete(system: str, context: str, user: str, *, max_tokens: int = 600) -> str:
    """One cached call. `context` (master CV + JD) is cache-pinned across bullets."""
    client = _client()
    resp = client.messages.create(
        model=MODEL,
        max_tokens=max_tokens,
        system=[
            {"type": "text", "text": system},
            # Cache the bulky, reused context so per-bullet calls are cheap.
            {"type": "text", "text": context, "cache_control": {"type": "ephemeral"}},
        ],
        messages=[{"role": "user", "content": user}],
    )
    return "".join(b.text for b in resp.content if b.type == "text").strip()


def rewrite_bullet(original: str, jd_text: str, master_cv: str) -> str:
    """Rewrite one CV bullet to XYZ form, tailored to the JD."""
    ctx = f"=== MASTER CV ===\n{master_cv}\n\n=== TARGET JOB DESCRIPTION ===\n{jd_text}"
    return _complete(_XYZ_SYSTEM, ctx, f"Rewrite this bullet:\n{original}")


# Words a bullet must never END on after a trim (would read as cut off).
_TRAIL_STOP = {
    "and", "or", "the", "a", "an", "to", "of", "for", "in", "on", "with", "by",
    "via", "at", "as", "that", "which", "while", "through", "from", "into", "per",
    "across", "using", "including", "&",
}


def _hard_trim(text: str, max_chars: int) -> str:
    """Last-resort deterministic shortener. Trims at a word boundary AND drops any
    trailing connector words so the bullet never ends on 'and', 'for', 'via', etc."""
    if len(text) <= max_chars:
        return text
    words = text[:max_chars].rsplit(" ", 1)[0].split()
    while words and words[-1].lower().strip(",;:-—&") in _TRAIL_STOP:
        words.pop()
    return " ".join(words).rstrip(" ,;:-—")


# The model must NEVER leak a clarification/refusal into a CV bullet. If a
# rewrite looks like meta-text (a question, an apology, "please share the JD"…),
# we discard it and keep the real bullet instead.
_META_RE = re.compile(
    r"please (share|provide|let me know)|could you (please )?(share|provide)"
    r"|i (can|could|will|would|need|am unable|cannot|can't|don't have|do not have)"
    r"|as an ai|i'?m sorry|i apologi[sz]e|happy to (help|assist)"
    r"|the (target |actual )?(job description|cv bullet)|provide the"
    r"|here('| i)s the rewritten|sure[,!]? ",
    re.IGNORECASE,
)


def _looks_like_meta(text: str) -> bool:
    t = (text or "").strip()
    return (not t) or t.endswith("?") or bool(_META_RE.search(t))


def _safe_rewrite(candidate: str, fallback: str) -> str:
    """Return the model's rewrite, or the fallback if it looks like meta/refusal."""
    return fallback if _looks_like_meta(candidate) else candidate.strip()


def _guidance_clause(guidance: str | None) -> str:
    if not guidance:
        return ""
    return ("\n\nAlso apply this recruiter feedback where it concerns wording or "
            "emphasis (IGNORE any advice about dates, graduation year, eligibility, "
            f"or whether to apply):\n{guidance}")


def rewrite_bullet_fitted(
    original: str, jd_text: str, master_cv: str,
    width_pt: float, font_name: str, size_pt: float,
    *, max_lines: int | None = None, max_tries: int = 3, guidance: str | None = None,
) -> str:
    """Rewrite a bullet to fill clean lines AND never grow taller than allowed.

    The target line count is the original bullet's (capped at 2), further capped
    by `max_lines` when the one-page backstop forces compression. The rewrite is
    HARD-bounded: if the model won't fit it, we trim to the line budget. This
    guarantees the tailored CV is never taller than the master -> stays one page.
    """
    from . import fit

    orig_n = min(max(fit.wrap_info(original, width_pt, font_name, size_pt)[0], 1), 2)
    target = orig_n if max_lines is None else min(orig_n, max_lines)
    cpl = fit.chars_per_line(width_pt, font_name, size_pt)
    lo, hi = fit.char_budget(target, cpl)
    ctx = f"=== MASTER CV ===\n{master_cv}\n\n=== TARGET JOB DESCRIPTION ===\n{jd_text}"

    if guidance:
        text = _safe_rewrite(
            _complete(_XYZ_SYSTEM, ctx,
                      f"Rewrite this bullet.{_guidance_clause(guidance)}\n\nBullet: {original}"),
            original)
    else:
        text = _safe_rewrite(rewrite_bullet(original, jd_text, master_cv), original)
    n, frac = fit.wrap_info(text, width_pt, font_name, size_pt)
    tries = 0
    while (n > target or not fit.is_clean(n, frac, target)) and tries < max_tries:
        if target == 1:
            user = (
                f"Rewrite this CV bullet as ONE single line, at most {hi} characters. "
                f"Lead with a strong past-tense verb and keep the single most important "
                f"number/metric. You MAY drop the 'by doing...' clause to fit. It must "
                f"NOT wrap onto a second line.{_guidance_clause(guidance)}\n\nBullet: {text}"
            )
        else:
            user = (
                f"Rewrite this CV bullet to between {lo} and {hi} characters so it fills "
                f"exactly two full lines (never a line and a quarter). Keep the XYZ format "
                f"and every number.{_guidance_clause(guidance)}\n\nBullet: {text}"
            )
        text = _safe_rewrite(_complete(_XYZ_SYSTEM, ctx, user), text)
        n, frac = fit.wrap_info(text, width_pt, font_name, size_pt)
        tries += 1

    # deterministic guarantee: never exceed the line budget
    if fit.wrap_info(text, width_pt, font_name, size_pt)[0] > target:
        text = _hard_trim(text, int(hi * 0.97))
    return text


_GROUP_SYSTEM = """You rewrite a numbered set of CV bullets that ALL belong to ONE
role/section, tailored to a target job. Hard rules:
- Return EXACTLY one rewritten bullet per input, in the same order, as a numbered
  list: "1.", "2.", ...
- Every bullet MUST be DISTINCT from its siblings: no two may share the same
  opening verb, the same metric/number, or the same phrasing. Preserve what makes
  each original bullet unique — do not let them collapse onto the same headline.
- Re-ANGLE each bullet to foreground what THIS job values (e.g. for a debt /
  restructuring / capital-structure role, lead with the credit, leverage or
  financing side; for research, lead with the thesis/analysis). Change emphasis
  and framing, not just wording — but stay strictly truthful; never invent.
- MATCH each bullet's original length: a bullet allowed two lines should FILL
  about two full lines with real substance (keep its specific detail) — do NOT
  over-compress it into a short one-liner. Only shorten when the limit forces it.
- Lead each bullet with a strong past-tense verb. Keep the key real number; never
  invent figures.
- Hit the per-bullet character range given — return COMPLETE sentences within it,
  never a clause cut short.
Return ONLY the numbered bullets, nothing else."""

_NUM_RE = re.compile(r"^\s*\d+[.)\-:]\s*(.+)$")


def _parse_numbered(raw: str, n: int) -> list[str]:
    out = [m.group(1).strip() for ln in raw.splitlines() if (m := _NUM_RE.match(ln))]
    return out if len(out) == n else []


def rewrite_group_fitted(
    originals: list[str], jd_text: str, master_cv: str,
    width_pt: float, font_name: str, size_pt: float,
    *, max_lines_each: int | None = None, guidance: str | None = None,
) -> list[str]:
    """Rewrite all bullets of one role together so they stay DISTINCT.

    One LLM call returns a tailored, fitted line per original. Falls back to
    independent per-bullet rewriting if the model returns the wrong count.
    """
    from . import fit

    if len(originals) == 1:
        return [rewrite_bullet_fitted(
            originals[0], jd_text, master_cv, width_pt, font_name, size_pt,
            max_lines=max_lines_each, guidance=guidance,
        )]

    cpl = fit.chars_per_line(width_pt, font_name, size_pt)
    targets = []
    for o in originals:
        t = min(max(fit.wrap_info(o, width_pt, font_name, size_pt)[0], 1), 2)
        targets.append(min(t, max_lines_each) if max_lines_each else t)

    req_lines = []
    for i, (o, t) in enumerate(zip(originals, targets), 1):
        lo, hi = fit.char_budget(t, cpl)
        req_lines.append(
            f"{i}. (fill {t} full line{'s' if t > 1 else ''}, {lo}-{hi} chars) {o}")
    ctx = f"=== MASTER CV ===\n{master_cv}\n\n=== TARGET JOB DESCRIPTION ===\n{jd_text}"
    user = ("Rewrite these bullets, keeping each one DISTINCT:\n" + "\n".join(req_lines)
            + _guidance_clause(guidance))

    raw = _complete(_GROUP_SYSTEM, ctx, user, max_tokens=900)
    parsed = _parse_numbered(raw, len(originals))
    if not parsed:  # model didn't comply -> safe fallback
        return [rewrite_bullet_fitted(
            o, jd_text, master_cv, width_pt, font_name, size_pt, max_lines=max_lines_each
        ) for o in originals]

    out = []
    for line, orig, t in zip(parsed, originals, targets):
        line = _safe_rewrite(line, orig)   # never let meta/refusal text become a bullet
        n, frac = fit.wrap_info(line, width_pt, font_name, size_pt)
        # Regenerate if too long (would overflow) OR too short (collapsed a
        # two-line bullet into one and lost detail it had room to keep).
        if not fit.is_clean(n, frac, t):
            line = rewrite_bullet_fitted(
                orig, jd_text, master_cv, width_pt, font_name, size_pt,
                max_lines=t, guidance=guidance,
            )
        out.append(line)
    return out


def _bullet_groups(doc) -> list[list[int]]:
    """Maximal runs of consecutive bullet paragraphs (= one role's bullets)."""
    groups: list[list[int]] = []
    cur: list[int] = []
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() and _is_list_item(p):
            cur.append(i)
        elif cur:
            groups.append(cur)
            cur = []
    if cur:
        groups.append(cur)
    return groups


def write_cover_paragraph(prompt: str, jd_text: str, master_cv: str) -> str:
    ctx = f"=== CANDIDATE BACKGROUND ===\n{master_cv}\n\n=== JOB DESCRIPTION ===\n{jd_text}"
    return _complete(_COVER_SYSTEM, ctx, prompt, max_tokens=500)


_COVER_LETTER_SYSTEM = """You write complete, concise cover letters for finance
internship applications. Structure: exactly three short paragraphs —
(1) the specific role and a genuine, specific reason for THIS firm (not generic);
(2) the candidate's most relevant evidence drawn ONLY from their CV, mapped to
what the role needs; (3) a brief, confident close.
Match the company's tone as inferred from the job description (formal/understated
vs. energetic). No clichés ("I am writing to express my interest"), no invented
experience, ~180-220 words. Return ONLY the letter body — no date, address, or
"Dear..."/"Sincerely" lines."""


def write_cover_letter(jd_text: str, master_cv: str, firm: str, name: str) -> str:
    """Generate a full, tone-matched cover letter body from the CV + JD."""
    ctx = f"=== CANDIDATE BACKGROUND ===\n{master_cv}\n\n=== JOB DESCRIPTION ===\n{jd_text}"
    user = f"Write the cover letter from {name} for the role at {firm}."
    return _complete(_COVER_LETTER_SYSTEM, ctx, user, max_tokens=700)


_REVIEW_SYSTEM = """You are an experienced finance recruiter reviewing a candidate's
tailored CV against a specific job description. Give concise, candid, actionable
feedback as markdown bullet points, grouped under three bold headings:
**What fits well** (2-3 points on the strongest alignments),
**Gaps vs this role** (the most important things the JD wants that the CV is
missing or weak on — include eligibility red flags like graduation year, location
or language if visible), and
**Specific fixes** (1-3 concrete edits: name the bullet/section and what to change).
Be direct and specific to THIS job. Never invent experience the CV doesn't show.
Keep the whole thing under ~180 words."""


_EXTRACT_SYSTEM = """You extract what an Applicant Tracking System and a recruiter
screen for from a job description. Return STRICT JSON ONLY (no prose, no code fence):
{"hard": [{"label": "Financial Modelling", "terms": ["financial model","three-statement","modelling"]}],
 "soft": [{"label": "Teamwork", "terms": ["team","collaborat"]}]}
Rules:
- "hard" = concrete technical skills, tools, methods, domain knowledge the role needs (aim 6-12).
- "soft" = competencies / attributes the JD explicitly asks for (aim 4-8).
- "terms" = 3-5 lowercase SHORT STEMS/single words an ATS would substring-match in a CV.
  Prefer stems over phrases: use "team","collaborat","lead" not "team environments";
  "organis","organiz","planning" not "organizational skills"; "valuation","comps",
  "dcf" not "company valuation analysis". Stems catch more real CV wording.
- Only include requirements the JD actually states. Be specific to THIS job."""


def _safe_json(raw: str) -> dict:
    raw = raw.strip()
    if "```" in raw:
        raw = raw.split("```")[1]
        raw = raw[4:] if raw.lower().startswith("json") else raw
    s, e = raw.find("{"), raw.rfind("}")
    return json.loads(raw[s:e + 1]) if s >= 0 else {}


def extract_requirements(jd_text: str) -> dict:
    """Parse the JD into hard-skill and competency requirements (one LLM call).

    The result is JD-only (no CV), so coverage can be re-checked for free with
    ats.check_coverage after any manual edit."""
    try:
        data = _safe_json(_complete(_EXTRACT_SYSTEM, jd_text,
                                    "Extract the requirements as JSON.", max_tokens=900))
        return {"hard": data.get("hard", []), "soft": data.get("soft", [])}
    except Exception:
        return {"hard": [], "soft": []}


def review_cv(cv_text: str, jd_text: str) -> str:
    """LLM critique of how well the tailored CV matches the JD, with fixes."""
    ctx = f"=== JOB DESCRIPTION ===\n{jd_text}\n\n=== TAILORED CV ===\n{cv_text}"
    return _complete(_REVIEW_SYSTEM, ctx, "Review this CV against the job description.",
                     max_tokens=600)


def save_cover_docx(text: str, out_path: str) -> str:
    """Save a cover-letter body as a simple, clean .docx (one para per block)."""
    from docx import Document

    doc = Document()
    for block in text.split("\n\n"):
        if block.strip():
            doc.add_paragraph(block.strip())
    doc.save(out_path)
    return out_path


from docx.oxml.ns import qn  # noqa: E402

# Verbs that mark an achievement bullet worth XYZ-rewriting. Education facts
# ("12 GCSEs…", "Relevant Modules:…", "Current Average:…") start otherwise and
# are left untouched so we never mangle grades.
_ACTION_VERBS = {
    "built", "led", "analyzed", "analysed", "valued", "traded", "pitched",
    "modeled", "modelled", "developed", "managed", "conducted", "created",
    "designed", "presented", "delivered", "achieved", "drove", "launched",
    "researched", "produced", "generated", "improved", "increased", "reduced",
    "negotiated", "advised", "structured", "executed", "spearheaded", "founded",
    "coordinated", "organized", "organised", "won", "ranked", "awarded",
    "completed", "prepared", "evaluated", "assessed", "forecast", "identified",
    "implemented", "established", "secured", "automated", "engineered",
}


def _is_list_item(paragraph) -> bool:
    """True if the paragraph is a real bullet/numbered list item or list style."""
    if paragraph._p.find(".//" + qn("w:numPr")) is not None:
        return True
    name = (paragraph.style.name or "") if paragraph.style else ""
    return "Bullet" in name or "List" in name


def _looks_like_achievement(text: str) -> bool:
    """Heuristic: does this bullet read like an accomplishment (vs. a fact list)?"""
    t = text.strip()
    if not t:
        return False
    first = t.split()[0].rstrip(":").lower()
    if first in _ACTION_VERBS:
        return True
    # a past-tense (-ed) verb leading the bullet is usually an achievement
    return first.isalpha() and first.endswith("ed") and len(first) > 3


def list_bullets(doc) -> list[tuple[int, str, bool]]:
    """Every bullet in the doc as (paragraph_index, text, looks_like_achievement)."""
    out: list[tuple[int, str, bool]] = []
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() and _is_list_item(p):
            out.append((i, p.text.strip(), _looks_like_achievement(p.text.strip())))
    return out


# Impactful CV metrics to bold: percentages, multiples, currency amounts, counts.
_METRIC_RE = re.compile(
    r'[£$€]\d[\d,.]*\s?[KMB]?n?\+?'   # £12,000+  $1B+  $10B  €5m
    r'|\d[\d,.]*\s?%'                 # 28.5%  83%
    r'|\d[\d,.]*x\b'                  # 3.5x  10x
    r'|\d[\d,.]*\+',                  # 50+  200+  30+
    re.IGNORECASE,
)


def _add_run(paragraph, text: str, name, size, bold: bool) -> None:
    run = paragraph.add_run(text)
    if name:
        run.font.name = name
    if size:
        run.font.size = size
    if bold:
        run.font.bold = True


def _replace_paragraph_text(paragraph, new_text: str) -> None:
    """Replace a bullet's text, keeping its font, and BOLD the impactful metrics
    (percentages, multiples, currency amounts, counts) like a polished CV does."""
    base = paragraph.runs[0] if paragraph.runs else None
    name = base.font.name if base else None
    size = base.font.size if base else None
    for r in list(paragraph.runs):           # clear existing runs
        r._element.getparent().remove(r._element)

    pos = 0
    for m in _METRIC_RE.finditer(new_text):
        if m.start() > pos:
            _add_run(paragraph, new_text[pos:m.start()], name, size, False)
        _add_run(paragraph, m.group(0), name, size, True)   # the metric -> bold
        pos = m.end()
    if pos < len(new_text) or not paragraph.runs:
        _add_run(paragraph, new_text[pos:], name, size, False)


@dataclass
class TailorResult:
    docx_path: str
    pdf_path: str | None
    bullets_rewritten: int
    page_count: int | None = None


def _page_count(pdf_path: str | None) -> int | None:
    if not pdf_path:
        return None
    try:
        from pypdf import PdfReader

        return len(PdfReader(pdf_path).pages)
    except Exception:
        return None


def tailor_cv(
    master_path: str,
    jd_text: str,
    out_docx: str,
    *,
    select: set[int] | None = None,
    to_pdf: bool = True,
    one_page: bool = True,
    guidance: str | None = None,
    keep_terms: list[str] | None = None,
) -> TailorResult:
    """Produce a tailored CV at `out_docx` with the master's exact layout.

    Objective order: (1) pack in as much of the JD as possible, (2) retain the
    original bullets' detail, (3) keep it to ONE page. Bullets are first written
    to fill their original line count (so detail isn't lost). If that overflows,
    only the FEWEST, LEAST job-relevant bullets are compressed — ranked by how
    few `keep_terms` (JD keywords) they contain — until it lands on one page, so
    the most JD-relevant detail survives.
    """
    from docx import Document

    from . import fit

    doc = Document(master_path)
    master_text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    width_pt, font_name, size_pt = _bullet_geometry(doc)

    bullets = list_bullets(doc)
    if select is None:
        select = {i for i, _, achievement in bullets if achievement}
    idx_text = {i: t for i, t, _ in bullets}
    groups = _bullet_groups(doc)

    # Rewrite each role's bullets together so siblings stay distinct.
    count = 0
    for grp in groups:
        gsel = [i for i in grp if i in select and idx_text.get(i)]
        if not gsel:
            continue
        news = rewrite_group_fitted(
            [idx_text[i] for i in gsel], jd_text, master_text,
            width_pt, font_name, size_pt, guidance=guidance,
        )
        for i, new in zip(gsel, news):
            _replace_paragraph_text(doc.paragraphs[i], new)
            count += 1

    doc.save(out_docx)
    pdf_path = export_pdf(out_docx) if to_pdf else None
    pages = _page_count(pdf_path)

    # One-page backstop. Ground truth is the REAL rendered page count (LibreOffice
    # + pypdf), never our width estimate — so it works even when the host's fonts
    # differ from Windows. Each round shortens the longest, least-job-relevant
    # bullet not yet shortened, re-renders, and stops the instant it fits one page.
    if one_page and pdf_path and pages and pages > 1:
        kt = [t.lower() for t in (keep_terms or [])]
        done: set[int] = set()
        for _ in range(len(select) + 3):
            doc = Document(out_docx)
            cands = []
            for idx in select:
                txt = doc.paragraphs[idx].text
                if idx in done or len(txt) < 55:   # too short to usefully shrink
                    continue
                hits = sum(1 for kw in kt if kw in txt.lower())
                cands.append((hits, -len(txt), idx))  # least relevant, then longest
            if not cands:
                break
            cands.sort()
            idx = cands[0][2]
            new = rewrite_bullet_fitted(
                idx_text[idx], jd_text, master_text, width_pt, font_name, size_pt,
                max_lines=1, guidance=guidance,
            )
            _replace_paragraph_text(doc.paragraphs[idx], new)
            done.add(idx)
            doc.save(out_docx)
            pdf_path = export_pdf(out_docx)
            pages = _page_count(pdf_path)
            if pages == 1:
                break

    return TailorResult(
        docx_path=out_docx, pdf_path=pdf_path, bullets_rewritten=count, page_count=pages
    )


_EMU_PER_PT = 12700


def apply_bullets(
    master_path: str, edits: dict[int, str], out_docx: str, *, to_pdf: bool = True
) -> TailorResult:
    """Rebuild the CV from the master, replacing chosen paragraphs with the user's
    own edited text. No LLM call — instant, for manual tweaks in the UI."""
    from docx import Document

    doc = Document(master_path)
    for idx, text in edits.items():
        if 0 <= idx < len(doc.paragraphs) and text.strip():
            _replace_paragraph_text(doc.paragraphs[idx], text.strip())
    doc.save(out_docx)
    pdf_path = export_pdf(out_docx) if to_pdf else None
    return TailorResult(
        docx_path=out_docx, pdf_path=pdf_path,
        bullets_rewritten=len(edits), page_count=_page_count(pdf_path),
    )


def _bullet_geometry(doc) -> tuple[float, str, float]:
    """Usable line width (pt), font name and size for bullet text in this doc."""
    sec = doc.sections[0]
    text_width_pt = (sec.page_width - sec.left_margin - sec.right_margin) / _EMU_PER_PT
    width_pt = text_width_pt - 24.0  # typical bullet hanging indent + glyph/space

    normal = doc.styles["Normal"].font
    name = normal.name or "Calibri"
    size = normal.size.pt if normal.size else 11.0
    # prefer an actual bullet paragraph's font/indent if it overrides Normal
    for para in doc.paragraphs:
        if para.style and "Bullet" in (para.style.name or "") and para.runs:
            f = para.runs[0].font
            name = f.name or name
            size = f.size.pt if f.size else size
            pf = para.paragraph_format
            if pf.left_indent:
                width_pt = text_width_pt - pf.left_indent.pt - 6.0
            break
    return (max(width_pt, 100.0), name, size)


_SOFFICE_PATHS = [
    r"C:\Program Files\LibreOffice\program\soffice.exe",
    r"C:\Program Files (x86)\LibreOffice\program\soffice.exe",
    "/usr/bin/soffice", "/snap/bin/libreoffice",
]


def _find_soffice() -> str | None:
    """Locate LibreOffice even when it isn't on PATH (the installer often skips it)."""
    found = shutil.which("soffice") or shutil.which("soffice.exe")
    if found:
        return found
    for path in _SOFFICE_PATHS:
        if os.path.exists(path):
            return path
    return None


def pdf_backend() -> str | None:
    """Which docx->pdf converter is available: 'word', 'libreoffice', or None."""
    import importlib.util

    if importlib.util.find_spec("docx2pdf"):
        return "word"
    if _find_soffice():
        return "libreoffice"
    return None


def export_pdf(docx_path: str) -> str | None:
    """Convert .docx -> .pdf preserving layout. Tries MS Word, then LibreOffice.

    Returns the pdf path, or None if no converter is installed (in which case the
    caller keeps the .docx and tells the user how to enable PDF export).
    """
    pdf = os.path.splitext(docx_path)[0] + ".pdf"
    backend = pdf_backend()
    if backend == "word":
        from docx2pdf import convert

        convert(docx_path, pdf)
    elif backend == "libreoffice":
        soffice = _find_soffice()
        out_dir = os.path.dirname(os.path.abspath(docx_path)) or "."
        subprocess.run(
            [soffice, "--headless", "--convert-to", "pdf", "--outdir", out_dir, docx_path],
            check=True, timeout=120,
        )
    else:
        return None
    return pdf if os.path.exists(pdf) else None
