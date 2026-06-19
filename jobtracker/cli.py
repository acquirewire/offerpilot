"""Command-line entrypoints. Run: python -m jobtracker <command> [args].

  init-db                       create the SQLite schema
  poll        --config c.yaml   run the Drop Tracker loop (Ctrl-C to stop)
  scan        --cv f --jd f     Mock ATS score of a CV against a JD
  tailor      --master f --jd f --out f --name "First Last" --firm "Goldman Sachs"
  autofill    --url U --profile p.yaml
  export      [--db d] [--out x.xlsx]
"""
from __future__ import annotations

import argparse
import asyncio
import sys


def _read(path: str) -> str:
    with open(path, "r", encoding="utf-8") as fh:
        return fh.read()


def cmd_init_db(args) -> None:
    from . import db

    conn = db.connect(args.db)
    db.init_db(conn)
    print(f"initialized {args.db}")


def cmd_poll(args) -> None:
    from .monitor import run

    try:
        asyncio.run(run(args.config, prime=getattr(args, "prime", False)))
    except KeyboardInterrupt:
        print("\nstopped.")


def cmd_scan(args) -> None:
    from .ats import scan

    result = scan(_read(args.cv), _read(args.jd))
    print(f"Match rate: {result.match_pct * 100:.1f}%")
    print(f"Matched   : {', '.join(result.matched) or '-'}")
    print(result.flag)
    sys.exit(0 if not result.missing else 1)


def cmd_tailor(args) -> None:
    from .ats import scan
    from .naming import output_filename
    from .tailor import tailor_cv

    jd = _read(args.jd)
    out = args.out or output_filename(args.name, args.firm)
    res = tailor_cv(args.master, jd, out.replace(".pdf", ".docx"))
    print(f"rewrote {res.bullets_rewritten} bullets -> {res.docx_path}")
    if res.pdf_path:
        print(f"pdf -> {res.pdf_path}")
    # Re-score the tailored doc so you see the lift.
    from docx import Document

    cv_text = "\n".join(p.text for p in Document(res.docx_path).paragraphs)
    print("post-tailor ATS:", scan(cv_text, jd).flag)


def _docx_text(path: str) -> str:
    """Plain text of a .docx (or a .txt) for ATS scoring."""
    if path.lower().endswith(".docx"):
        from docx import Document

        return "\n".join(p.text for p in Document(path).paragraphs)
    return _read(path)


def cmd_apply(args) -> None:
    """Chain the whole application: limit-check -> scan -> tailor -> log -> autofill."""
    import os

    from . import db
    from .ats import scan
    from .naming import firm_slug, output_filename

    jd = _read(args.jd)
    conn = db.connect(args.db)
    db.init_db(conn)
    firm_id = db.get_or_create_firm(conn, args.firm, firm_slug(args.firm), max_apps=args.max_apps)

    # --- Rule Engine: warn BEFORE spending a slot ---
    active, mx, at_limit = db.check_firm_limit(conn, firm_id)
    print(f"\n[{args.firm}] {active}/{mx} applications used.")
    if at_limit and not args.force:
        print(f"  STOP: you're at the {args.firm} limit ({mx}). Re-run with --force to override.")
        return

    # --- Tailor CV + cover letter (if API key present) or fall back ---
    cv_file = args.master
    cover_file = None
    out_docx = output_filename(args.name, args.firm, ext="docx")
    if os.environ.get("ANTHROPIC_API_KEY") and not args.no_tailor:
        from .tailor import save_cover_docx, tailor_cv, write_cover_letter

        print("Tailoring CV to the job description...")
        res = tailor_cv(args.master, jd, out_docx)
        cv_file = res.pdf_path or res.docx_path
        cv_text = _docx_text(res.docx_path)
        print(f"  rewrote {res.bullets_rewritten} bullets -> {cv_file}")

        print("Writing cover letter...")
        cover_out = output_filename(args.name, args.firm, kind="Cover", ext="docx")
        save_cover_docx(write_cover_letter(jd, cv_text, args.firm, args.name), cover_out)
        cover_file = cover_out
        print(f"  -> {cover_file}")
    else:
        cv_text = _docx_text(args.master)
        if not args.no_tailor:
            print("  (no ANTHROPIC_API_KEY — skipping tailoring, using master CV as-is)")

    # --- ATS score ---
    result = scan(cv_text, jd)
    print(f"ATS match: {result.match_pct*100:.0f}%  | {result.flag}")

    # --- Log it to the tracker ---
    posting_id = db.get_or_create_posting(
        conn, firm_id, args.role, cycle=args.cycle, apply_url=args.url
    )
    app_id = db.log_application(
        conn, posting_id, firm_id,
        cv_file=os.path.basename(cv_file),
        cover_file=os.path.basename(cover_file) if cover_file else None,
        ats_match_pct=result.match_pct,
    )
    print(f"Logged application #{app_id} to the tracker.")

    # --- Autofill (optional) ---
    if args.url and args.profile:
        import yaml

        from .autofill import Profile, autofill

        data = yaml.safe_load(_read(args.profile))
        data.setdefault("cv_path", cv_file)
        print("Opening the application form for review (you click Submit)...")
        autofill(Profile(**data), args.url)
    elif args.url:
        print(f"\nNext: apply here -> {args.url}")


def cmd_cover(args) -> None:
    """Generate a tone-matched cover letter from your CV + the job description."""
    from .naming import output_filename
    from .tailor import save_cover_docx, write_cover_letter

    jd = _read(args.jd)
    master = _docx_text(args.master) if args.master else ""
    text = write_cover_letter(jd, master, args.firm, args.name)
    out = args.out or output_filename(args.name, args.firm, kind="Cover", ext="docx")
    save_cover_docx(text, out)
    print(f"cover letter saved -> {out}\n")
    print("--- preview " + "-" * 40)
    print(text)


def cmd_autofill(args) -> None:
    import yaml

    from .autofill import Profile, autofill

    data = yaml.safe_load(_read(args.profile))
    autofill(Profile(**data), args.url)


def cmd_export(args) -> None:
    from . import db
    from .export import export_xlsx

    conn = db.connect(args.db)
    path = export_xlsx(conn, args.out)
    print(f"exported -> {path}")


def main(argv: list[str] | None = None) -> None:
    p = argparse.ArgumentParser(prog="jobtracker")
    sub = p.add_subparsers(dest="cmd", required=True)

    s = sub.add_parser("init-db"); s.add_argument("--db", default="jobtracker.db"); s.set_defaults(fn=cmd_init_db)
    s = sub.add_parser("poll"); s.add_argument("--config", required=True)
    s.add_argument("--prime", action="store_true", help="one silent baseline pass (no alerts), then exit")
    s.set_defaults(fn=cmd_poll)
    s = sub.add_parser("scan"); s.add_argument("--cv", required=True); s.add_argument("--jd", required=True); s.set_defaults(fn=cmd_scan)
    s = sub.add_parser("tailor")
    s.add_argument("--master", required=True); s.add_argument("--jd", required=True)
    s.add_argument("--out"); s.add_argument("--name", required=True); s.add_argument("--firm", required=True)
    s.set_defaults(fn=cmd_tailor)
    s = sub.add_parser("apply", help="scan + tailor + log + (optional) autofill in one step")
    s.add_argument("--firm", required=True); s.add_argument("--role", required=True)
    s.add_argument("--master", required=True); s.add_argument("--jd", required=True)
    s.add_argument("--name", required=True)
    s.add_argument("--url"); s.add_argument("--profile"); s.add_argument("--cycle")
    s.add_argument("--db", default="jobtracker.db"); s.add_argument("--max-apps", type=int, default=3)
    s.add_argument("--no-tailor", action="store_true"); s.add_argument("--force", action="store_true")
    s.set_defaults(fn=cmd_apply)
    s = sub.add_parser("cover", help="generate a tone-matched cover letter (.docx)")
    s.add_argument("--jd", required=True); s.add_argument("--master")
    s.add_argument("--name", required=True); s.add_argument("--firm", required=True)
    s.add_argument("--out"); s.set_defaults(fn=cmd_cover)
    s = sub.add_parser("autofill"); s.add_argument("--url", required=True); s.add_argument("--profile", required=True); s.set_defaults(fn=cmd_autofill)
    s = sub.add_parser("export"); s.add_argument("--db", default="jobtracker.db"); s.add_argument("--out", default="job_tracker.xlsx"); s.set_defaults(fn=cmd_export)

    args = p.parse_args(argv)
    args.fn(args)


if __name__ == "__main__":
    main()
