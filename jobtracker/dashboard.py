"""Pipeline dashboard (Module 4).  Run:  streamlit run jobtracker/dashboard.py

A styled control center over the SQLite tracker:
  * headline metrics + Rule-Engine limit warnings
  * a colour-coded pipeline board with one-click stage advancement
  * the live "open drops" feed from the monitor, with apply links
  * a Log-Application form, so you never need a code snippet

Requires streamlit + pandas.
"""
from __future__ import annotations

import os
import sqlite3
import sys

# `streamlit run jobtracker/dashboard.py` puts THIS file's folder on sys.path,
# not the project root — so `import jobtracker.*` fails. Add the repo root.
sys.path.insert(0, os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import pandas as pd
import streamlit as st

try:  # load .env (local) so keys are available
    from dotenv import load_dotenv

    load_dotenv()
except ImportError:
    pass

# On Streamlit Community Cloud (and other hosts), secrets are provided via
# st.secrets, not a .env file — mirror them into env vars so all the code that
# reads os.environ (API key, admin login, Stripe) works online unchanged.
try:
    for _k, _v in st.secrets.items():
        os.environ.setdefault(_k, str(_v))
except Exception:
    pass

DB_PATH = os.environ.get("JOBTRACKER_DB", "jobtracker.db")
STAGES = ["Applied", "OA", "HireVue", "Superday", "Offer", "Rejected"]
STAGE_COLOR = {
    "Applied": "#3b82f6", "OA": "#8b5cf6", "HireVue": "#f59e0b",
    "Superday": "#ec4899", "Offer": "#10b981", "Rejected": "#94a3b8",
}

st.set_page_config(page_title="Job Pipeline", page_icon="💼", layout="wide")

# --- styling -----------------------------------------------------------------
st.markdown(
    """
    <style>
      @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;500;600;700&display=swap');
      html, body, [class*="css"] { font-family: 'Inter', sans-serif; }
      #MainMenu, footer, header { visibility: hidden; }
      .block-container { padding-top: 2rem; max-width: 1400px; }
      .hero {
        background: linear-gradient(135deg, #0f172a 0%, #1e3a5f 100%);
        color: #fff; padding: 1.5rem 1.75rem; border-radius: 16px; margin-bottom: 1.5rem;
      }
      .hero h1 { margin: 0; font-size: 1.6rem; font-weight: 700; }
      .hero p  { margin: .25rem 0 0; opacity: .75; font-size: .9rem; }
      .metric {
        background: #fff; border: 1px solid #e2e8f0; border-radius: 14px;
        padding: 1rem 1.25rem; box-shadow: 0 1px 3px rgba(0,0,0,.04);
      }
      .metric .v { font-size: 1.9rem; font-weight: 700; color: #0f172a; line-height: 1; }
      .metric .l { font-size: .78rem; color: #64748b; text-transform: uppercase;
                   letter-spacing: .04em; margin-top: .35rem; }
      .card {
        background: #fff; border-radius: 12px; padding: .8rem .9rem; margin-bottom: .7rem;
        box-shadow: 0 1px 3px rgba(0,0,0,.06); border-left: 4px solid #cbd5e1;
        color: #1e293b;   /* default dark text so nothing inherits the theme's white */
      }
      .card mark { background: #fde68a; color: #1e293b; padding: 0 2px; border-radius: 3px; }
      .card .firm { font-weight: 600; color: #0f172a; font-size: .92rem; }
      .card .role { color: #475569; font-size: .82rem; margin-top: .15rem; }
      .pill { display:inline-block; padding:.1rem .5rem; border-radius:999px;
              font-size:.7rem; font-weight:600; color:#fff; }
      .stagehead { font-weight:600; font-size:.85rem; color:#334155; margin-bottom:.5rem;
                   text-transform:uppercase; letter-spacing:.03em; }
    </style>
    """,
    unsafe_allow_html=True,
)


def conn() -> sqlite3.Connection:
    c = sqlite3.connect(DB_PATH)
    c.row_factory = sqlite3.Row
    return c


@st.cache_data(ttl=15)
def q(sql: str, params: tuple = ()) -> pd.DataFrame:
    with conn() as c:
        return pd.read_sql_query(sql, c, params=params)


def refresh():
    st.cache_data.clear()
    st.rerun()


# --- accounts: sign-up / login / tiers / billing -----------------------------
import jobtracker.accounts as acc
import jobtracker.billing as billing

acc.init()
acc.ensure_admin()

# Ensure the application-tracker tables exist. Locally these are made by the
# `init-db` command, but on a fresh host (Streamlit Cloud) nothing has run yet,
# so create the schema on first load to avoid "no such table" errors.
try:
    import jobtracker.db as _jdb

    _jc = _jdb.connect(DB_PATH)
    _jdb.init_db(_jc)
    _jc.close()
except Exception:
    pass

# Stripe Checkout return: ?session_id=... -> confirm payment -> unlock Pro
_qp = st.query_params
if "session_id" in _qp:
    _paid, _em = billing.verify_session(_qp.get("session_id"))
    if _paid and _em:
        acc.set_tier(_em, "pro")
        st.session_state["just_upgraded"] = True
    st.query_params.clear()


def _auth_gate() -> dict:
    if st.session_state.get("user"):
        # refresh tier from DB so admin changes take effect live
        u = acc.get_user(st.session_state["user"]["email"]) or st.session_state["user"]
        st.session_state["user"] = u
        return u
    st.markdown("<div class='hero'><h1>✈️ OfferPilot</h1>"
                "<p>Sign in to tailor your applications.</p></div>", unsafe_allow_html=True)
    t_in, t_up = st.tabs(["Log in", "Create account"])
    with t_in:
        e = st.text_input("Email", key="li_e")
        p = st.text_input("Password", type="password", key="li_p")
        if st.button("Log in", type="primary"):
            u = acc.verify_login(e, p)
            if u:
                st.session_state["user"] = u
                st.rerun()
            else:
                st.error("Wrong email or password.")
    with t_up:
        e = st.text_input("Email", key="su_e")
        p = st.text_input("Password (6+ characters)", type="password", key="su_p")
        if st.button("Create free account", type="primary"):
            ok, msg = acc.create_user(e, p)
            if ok:
                st.session_state["user"] = acc.get_user(e)
                st.rerun()
            else:
                st.error(msg)
    st.stop()


user = _auth_gate()

# --- account sidebar ---------------------------------------------------------
with st.sidebar:
    st.markdown(f"**{user['email']}**")
    badge = {"admin": "🛠️ Admin", "pro": "⭐ Pro", "free": "Free"}[user["tier"]]
    st.caption(f"Plan: {badge}")
    if st.session_state.pop("just_upgraded", False):
        st.success("You're Pro! 🎉")
    if user["tier"] == "free":
        st.caption(f"{acc.remaining_free(user)} free tailors left this month")
        if billing.configured():
            if st.button("⭐ Upgrade to Pro — £9/mo"):
                st.session_state["checkout_url"] = billing.create_checkout_url(
                    user["email"], os.environ.get("APP_URL", "http://localhost:8501"))
            if st.session_state.get("checkout_url"):
                st.link_button("Continue to secure checkout →", st.session_state["checkout_url"])
        else:
            st.caption("💳 Payments not set up yet.")
    if st.button("Log out"):
        st.session_state.pop("user", None)
        st.rerun()

# --- hero --------------------------------------------------------------------
st.markdown(
    "<div class='hero'><h1>💼 Application Pipeline</h1>"
    "<p>Live tracker across your target firms — monitored, scored, logged.</p></div>",
    unsafe_allow_html=True,
)

apps = q("""
    SELECT a.id, f.name AS firm, p.title AS role, a.stage, a.ats_match_pct, a.applied_at
    FROM application a JOIN firm f ON f.id=a.firm_id JOIN posting p ON p.id=a.posting_id
""")
open_posts = q("SELECT COUNT(*) n FROM posting WHERE status='OPEN'")
limits = q("SELECT name, active_apps, max_apps, at_limit FROM firm_load")

# --- metrics -----------------------------------------------------------------
total = len(apps)
active = len(apps[~apps["stage"].isin(["Rejected"])]) if total else 0
offers = len(apps[apps["stage"] == "Offer"]) if total else 0
n_open = int(open_posts["n"].iloc[0]) if not open_posts.empty else 0
cards = [("Applications", total), ("Active", active), ("Offers", offers), ("Open drops", n_open)]
for col, (label, val) in zip(st.columns(4), cards):
    col.markdown(f"<div class='metric'><div class='v'>{val}</div><div class='l'>{label}</div></div>",
                 unsafe_allow_html=True)

# --- rule-engine warnings ----------------------------------------------------
at = limits[limits["at_limit"] == 1] if not limits.empty else limits
if not at.empty:
    st.write("")
    for _, r in at.iterrows():
        st.warning(f"⚠️ **{r['name']}** — at application limit ({r['active_apps']}/{r['max_apps']})")

st.write("")
_labels = ["📋 Pipeline", "🔔 Open Drops", "➕ Log Application", "✍️ Tailor CV"]
if acc.is_admin(user):
    _labels.append("🛠️ Admin")
_tabs = st.tabs(_labels)
tab_pipe, tab_drops, tab_log, tab_tailor = _tabs[0], _tabs[1], _tabs[2], _tabs[3]
tab_admin = _tabs[4] if acc.is_admin(user) else None

# --- pipeline board ----------------------------------------------------------
with tab_pipe:
    if total == 0:
        st.info("No applications yet. Use **Log Application** or the `apply` command.")
    else:
        for col, stage in zip(st.columns(len(STAGES)), STAGES):
            color = STAGE_COLOR[stage]
            sub = apps[apps["stage"] == stage]
            col.markdown(f"<div class='stagehead'>{stage} · {len(sub)}</div>", unsafe_allow_html=True)
            for _, row in sub.iterrows():
                pct = f"<span class='pill' style='background:{color}'>{row['ats_match_pct']*100:.0f}%</span>" \
                      if pd.notna(row["ats_match_pct"]) else ""
                col.markdown(
                    f"<div class='card' style='border-left-color:{color}'>"
                    f"<div class='firm'>{row['firm']} {pct}</div>"
                    f"<div class='role'>{row['role']}</div></div>",
                    unsafe_allow_html=True,
                )
                nxt = STAGES[min(STAGES.index(stage) + 1, len(STAGES) - 1)]
                if stage not in ("Offer", "Rejected") and col.button(f"→ {nxt}", key=f"adv{row['id']}"):
                    with conn() as c:
                        from datetime import datetime, timezone
                        now = datetime.now(timezone.utc).isoformat(timespec="seconds")
                        c.execute("UPDATE application SET stage=? WHERE id=?", (nxt, row["id"]))
                        c.execute("INSERT INTO stage_event (application_id, stage, occurred_at) VALUES (?,?,?)",
                                  (row["id"], nxt, now))
                        c.commit()
                    refresh()

# --- open drops feed ---------------------------------------------------------
with tab_drops:
    drops = q("""
        SELECT f.name AS firm, p.title AS role, p.location, p.cycle, p.apply_url, p.last_seen
        FROM posting p JOIN firm f ON f.id=p.firm_id
        WHERE p.status='OPEN' ORDER BY p.last_seen DESC LIMIT 200
    """)
    if drops.empty:
        st.info("No open postings recorded yet. Start the monitor: `python -m jobtracker poll`.")
    else:
        st.caption(f"{len(drops)} open postings (most recent first)")
        for _, r in drops.iterrows():
            loc = f" · {r['location']}" if r["location"] else ""
            cyc = f" · {r['cycle']}" if r["cycle"] else ""
            link = f"<a href='{r['apply_url']}' target='_blank'>Apply ↗</a>" if r["apply_url"] else ""
            st.markdown(
                f"<div class='card'><div class='firm'>{r['firm']} {link}</div>"
                f"<div class='role'>{r['role']}{loc}{cyc}</div></div>",
                unsafe_allow_html=True,
            )

# --- log application form ----------------------------------------------------
with tab_log:
    firms = q("SELECT id, name FROM firm ORDER BY name")
    if firms.empty:
        st.info("No firms yet — they appear once the monitor runs or you apply once.")
    else:
        with st.form("logform", clear_on_submit=True):
            fsel = st.selectbox("Firm", firms["name"].tolist())
            role = st.text_input("Role", placeholder="2027 Summer Analyst, IBD")
            c1, c2 = st.columns(2)
            cv = c1.text_input("CV file", placeholder="Henry_Smith_GoldmanSachs.pdf")
            match = c2.slider("ATS match %", 0, 100, 75)
            if st.form_submit_button("Log application") and role:
                import jobtracker.db as jdb
                with conn() as c:
                    fid = c.execute("SELECT id FROM firm WHERE name=?", (fsel,)).fetchone()["id"]
                dbc = jdb.connect(DB_PATH)
                pid = jdb.get_or_create_posting(dbc, fid, role)
                jdb.log_application(dbc, pid, fid, cv_file=cv or None, ats_match_pct=match / 100)
                st.success(f"Logged {fsel} — {role}")
                refresh()

# --- tailor CV + cover letter ------------------------------------------------
with tab_tailor:
    import tempfile

    st.caption("Upload your master CV, paste the job description, and get a tailored "
               "CV + cover letter to download. Uses Claude (needs ANTHROPIC_API_KEY).")
    if not os.environ.get("ANTHROPIC_API_KEY"):
        st.error("No ANTHROPIC_API_KEY found in your .env — tailoring can't run without it.")

    up = st.file_uploader("Master CV — your Word file (.docx, not a PDF)", type=["docx"])

    # Detect the CV's bullets and let the user pick which to rewrite. Achievements
    # are pre-selected; grades/modules/skills are left out so they're never mangled.
    sel_idx: set[int] = set()
    if up:
        from io import BytesIO

        import jobtracker.tailor as _tl
        from docx import Document as _Doc

        _bullets = _tl.list_bullets(_Doc(BytesIO(up.getvalue())))
        if _bullets:
            _labels = {f"{i} · {t[:75]}": i for i, t, _ in _bullets}
            _default = [lbl for lbl, i in _labels.items()
                        if next(a for j, _, a in _bullets if j == i)]
            st.markdown("**Which bullets should Claude rewrite?** "
                        "(grades, modules and skills are left out on purpose)")
            _chosen = st.multiselect("Bullets to tailor", list(_labels), default=_default,
                                     label_visibility="collapsed")
            sel_idx = {_labels[c] for c in _chosen}
        else:
            st.warning("No bullet points detected in this file — is it the right CV?")

    jd_text = st.text_area("Job description", height=180, placeholder="Paste the full job description…")
    c1, c2 = st.columns(2)
    name = c1.text_input("Your full name", placeholder="Henry Smith")
    firm = c2.text_input("Firm", placeholder="Goldman Sachs")

    if st.button("✨ Tailor my CV + write cover letter", type="primary"):
        if not acc.can_tailor(user):
            st.warning(f"You've used your {acc.FREE_MONTHLY_LIMIT} free tailors this month. "
                       "Upgrade to Pro (sidebar) for unlimited tailoring.")
        elif not (up and jd_text.strip() and name and firm and sel_idx):
            st.warning("Upload a CV, pick at least one bullet, and fill in the job "
                       "description, name and firm first.")
        else:
            try:
                import jobtracker.tailor as tl
                from docx import Document

                from jobtracker.ats import jd_terms, scan
                from jobtracker.naming import output_filename

                with st.spinner("Tailoring + fitting to one page — this can take "
                                "2–3 minutes for a full CV. Hang tight…"):
                    tmpd = tempfile.mkdtemp()
                    master_path = os.path.join(tmpd, "master.docx")
                    with open(master_path, "wb") as fh:
                        fh.write(up.getbuffer())
                    before = [t for i, t, _ in tl.list_bullets(Document(master_path))
                              if i in sel_idx]
                    reqs = tl.extract_requirements(jd_text)
                    keep_terms = [t for grp in reqs.values() for it in grp
                                  for t in it.get("terms", [])]
                    cv_out = os.path.join(tmpd, output_filename(name, firm, ext="docx"))
                    res = tl.tailor_cv(master_path, jd_text, cv_out, select=sel_idx,
                                       keep_terms=keep_terms, to_pdf=True)
                    after_doc = Document(res.docx_path)
                    after_pairs = [(i, t) for i, t, _ in tl.list_bullets(after_doc) if i in sel_idx]
                    after = [t for _, t in after_pairs]
                    after_map = {i: t for i, t in after_pairs}
                    sel_order = [i for i, _ in after_pairs]
                    cv_text = "\n".join(p.text for p in after_doc.paragraphs)
                    cover = tl.write_cover_letter(jd_text, cv_text, firm, name)
                    cover_out = os.path.join(tmpd, output_filename(name, firm, kind="Cover", ext="docx"))
                    tl.save_cover_docx(cover, cover_out)
                    cover_pdf = tl.export_pdf(cover_out)  # None if no converter
                    ats = scan(cv_text, jd_text)
                    payload = {
                        "before": before, "after": after, "cover": cover,
                        "ats": f"{ats.match_pct*100:.0f}% — {ats.flag}",
                        "match_pct": ats.match_pct, "matched": ats.matched,
                        "missing": ats.missing, "jd_terms": jd_terms(jd_text),
                        "reqs": reqs, "jd": jd_text, "cv_text": cv_text,
                        "after_map": after_map, "sel_order": sel_order,
                        "master_bytes": up.getvalue(),
                        "cv_bytes": open(res.docx_path, "rb").read(),
                        "cv_name": os.path.basename(cv_out),
                        "cover_bytes": open(cover_out, "rb").read(),
                        "cover_name": os.path.basename(cover_out),
                        "cv_pdf": open(res.pdf_path, "rb").read() if res.pdf_path else None,
                        "cover_pdf": open(cover_pdf, "rb").read() if cover_pdf else None,
                        "pdf_name": output_filename(name, firm, ext="pdf"),
                        "cover_pdf_name": output_filename(name, firm, kind="Cover", ext="pdf"),
                        "pages": res.page_count,
                    }
                    st.session_state["tlr"] = payload
                    acc.record_use(user["email"])  # count against the free quota
                    st.session_state.pop("cv_review", None)  # clear old review
                    for _k in [k for k in st.session_state if k.startswith("edit_")]:
                        del st.session_state[_k]  # clear stale per-bullet edits
            except Exception as exc:  # noqa: BLE001 - surface API/credit errors to the user
                st.error(f"Tailoring failed: {exc}")

    r = st.session_state.get("tlr")
    if r:
        from jobtracker.ats import bullet_hits, check_coverage, highlight_html
        rep = check_coverage(r["cv_text"], r.get("reqs") or {"hard": [], "soft": []})
        pages = r.get("pages")
        pg = f" · {pages} page{'s' if pages and pages != 1 else ''}" if pages else ""
        st.success(f"Done — {rep.overall_pct*100:.0f}% fit to this job{pg}")
        if pages and pages > 1:
            st.warning("Still over one page — your master CV may have too much content "
                       "to fit even when compressed. Consider trimming a bullet or two.")
        prev = st.session_state.get("tlr_prev")
        if prev:
            old = check_coverage(prev["cv_text"], prev.get("reqs") or {}).overall_pct * 100
            new = rep.overall_pct * 100
            if new < old - 1:
                st.warning(f"⚠️ That improvement pass moved the fit {old:.0f}% → {new:.0f}%. "
                           "Re-angling for the role can trade off keyword coverage — use Undo "
                           "if you preferred the previous version.")
            if st.button("↩ Undo last improvement"):
                st.session_state["tlr"] = prev
                st.session_state.pop("tlr_prev", None)
                for _k in [k for k in st.session_state if k.startswith("edit_")]:
                    del st.session_state[_k]
                st.rerun()
        if r.get("cv_pdf"):
            p1, p2 = st.columns(2)
            p1.download_button("⬇ Download CV (PDF)", r["cv_pdf"], r["pdf_name"], type="primary")
            if r.get("cover_pdf"):
                p2.download_button("⬇ Download cover letter (PDF)", r["cover_pdf"], r["cover_pdf_name"], type="primary")
        else:
            st.info("📄 PDF export is off — no converter installed. Install LibreOffice "
                    "(free) to download PDFs; until then, use the Word files below. "
                    "Ask Claude: \"install LibreOffice for PDF export\".")
        d1, d2 = st.columns(2)
        d1.download_button("⬇ Download tailored CV (Word)", r["cv_bytes"], r["cv_name"])
        d2.download_button("⬇ Download cover letter (Word)", r["cover_bytes"], r["cover_name"])

        # ---- ATS & Fit feedback (requirements parsed from THIS JD) --------------
        st.divider()
        st.subheader("📊 ATS & fit against this job")
        pct = rep.overall_pct
        bar = "#10b981" if pct >= 0.8 else "#f59e0b" if pct >= 0.55 else "#ef4444"
        nreq = len(rep.hard) + len(rep.soft)
        st.markdown(
            f"<div style='font-size:2rem;font-weight:700;color:{bar}'>{pct*100:.0f}% overall fit</div>"
            f"<div style='background:#e2e8f0;border-radius:8px;height:10px;margin:.3rem 0 .2rem'>"
            f"<div style='width:{pct*100:.0f}%;background:{bar};height:10px;border-radius:8px'></div></div>"
            f"<div class='role' style='margin-bottom:1rem'>Hard skills {rep.hard_pct*100:.0f}% · "
            f"Competencies {rep.soft_pct*100:.0f}% · {nreq} requirements read from this job description"
            "</div>",
            unsafe_allow_html=True,
        )

        def _req_block(title, items):
            if not items:
                return
            st.markdown(f"**{title}**")
            chips = "".join(
                f"<span class='pill' style='background:"
                f"{'#10b981' if x.covered else '#ef4444'};margin:2px'>"
                f"{'✓' if x.covered else '✗'} {x.label}</span>" for x in items)
            st.markdown(chips, unsafe_allow_html=True)

        _req_block(f"Hard skills ({sum(x.covered for x in rep.hard)}/{len(rep.hard)})", rep.hard)
        _req_block(f"Competencies ({sum(x.covered for x in rep.soft)}/{len(rep.soft)})", rep.soft)

        terms = sorted({t for grp in (r.get("reqs") or {}).values()
                        for it in grp for t in it.get("terms", [])}, key=len, reverse=True)
        st.markdown("**Your bullets — JD keywords highlighted** "
                    "<span class='role'>(amber = no JD keyword, consider tying it to the role)</span>",
                    unsafe_allow_html=True)
        for a in r["after"]:
            hits = bullet_hits(a, terms)
            edge = "#10b981" if hits else "#f59e0b"
            st.markdown(
                f"<div class='card' style='border-left-color:{edge}'>"
                f"<span class='pill' style='background:{edge}'>{hits} JD hit{'s' if hits != 1 else ''}</span> "
                f"<span style='font-size:.86rem'>{highlight_html(a, terms)}</span></div>",
                unsafe_allow_html=True,
            )

        # AI improvement review + one-click apply
        bcol1, bcol2 = st.columns(2)
        if bcol1.button("💡 Get AI improvement tips"):
            import jobtracker.tailor as tl
            with st.spinner("Reviewing your CV against the job…"):
                st.session_state["cv_review"] = tl.review_cv(r["cv_text"], r["jd"])
        if bcol2.button("⚡ Apply these improvements", type="primary"):
            import tempfile

            import jobtracker.tailor as tl
            from docx import Document

            from jobtracker.ats import check_coverage, scan
            with st.spinner("Rewriting your bullets with the improvements — ~1 minute…"):
                st.session_state["tlr_prev"] = dict(r)  # snapshot so Undo can restore
                review = st.session_state.get("cv_review") or tl.review_cv(r["cv_text"], r["jd"])
                st.session_state["cv_review"] = review
                # Target the actual gaps + keep what's already covered.
                rep0 = check_coverage(r["cv_text"], r.get("reqs") or {})
                miss = [x for x in rep0.hard + rep0.soft if not x.covered]
                keyword_note = ("\n\nCRITICAL: where genuinely truthful to the candidate's "
                                "experience, work in these missing job requirements (use the "
                                "keyword): " + "; ".join(f"{x.label} ({x.terms[0]})" for x in miss)
                                + ". Do NOT drop keywords already present and do NOT fabricate.") if miss else ""
                tmpd = tempfile.mkdtemp()
                mp = os.path.join(tmpd, "master.docx")
                with open(mp, "wb") as fh:
                    fh.write(r["master_bytes"])
                out_docx = os.path.join(tmpd, r["cv_name"])
                keep_terms = [t for grp in (r.get("reqs") or {}).values()
                              for it in grp for t in it.get("terms", [])]
                res = tl.tailor_cv(mp, r["jd"], out_docx, select=set(r["sel_order"]),
                                   guidance=review + keyword_note, keep_terms=keep_terms, to_pdf=True)
                adoc = Document(res.docx_path)
                pairs = [(i, t) for i, t, _ in tl.list_bullets(adoc) if i in set(r["sel_order"])]
                cv_text = "\n".join(p.text for p in adoc.paragraphs)
                ats = scan(cv_text, r["jd"])
                r.update({
                    "after": [t for _, t in pairs], "after_map": dict(pairs),
                    "cv_text": cv_text, "match_pct": ats.match_pct,
                    "matched": ats.matched, "missing": ats.missing,
                    "ats": f"{ats.match_pct*100:.0f}% — {ats.flag}",
                    "cv_bytes": open(res.docx_path, "rb").read(),
                    "cv_pdf": open(res.pdf_path, "rb").read() if res.pdf_path else None,
                    "pages": res.page_count,
                })
                st.session_state["tlr"] = r
                for _k in [k for k in st.session_state if k.startswith("edit_")]:
                    del st.session_state[_k]
            st.rerun()
        if st.session_state.get("cv_review"):
            st.markdown(st.session_state["cv_review"])

        # ---- edit bullets yourself + instant rebuild ---------------------------
        st.divider()
        st.subheader("✏️ Edit bullets & rebuild")
        st.caption("Tweak any wording yourself — e.g. lead a bullet with the debt / "
                   "capital-structure angle for this role — then rebuild. No AI, instant.")
        edits = {}
        for idx in r["sel_order"]:
            edits[idx] = st.text_area(
                f"bullet_{idx}", value=r["after_map"][idx],
                key=f"edit_{idx}", height=68, label_visibility="collapsed",
            )
        if st.button("🔁 Rebuild CV with my edits", type="primary"):
            import tempfile

            import jobtracker.tailor as tl
            from docx import Document

            from jobtracker.ats import scan
            with st.spinner("Rebuilding your CV…"):
                tmpd = tempfile.mkdtemp()
                mp = os.path.join(tmpd, "master.docx")
                with open(mp, "wb") as fh:
                    fh.write(r["master_bytes"])
                out_docx = os.path.join(tmpd, r["cv_name"])
                res = tl.apply_bullets(mp, edits, out_docx, to_pdf=True)
                cv_text = "\n".join(p.text for p in Document(res.docx_path).paragraphs)
                ats = scan(cv_text, r["jd"])
                r.update({
                    "after": [edits[i] for i in r["sel_order"]],
                    "after_map": dict(edits), "cv_text": cv_text,
                    "match_pct": ats.match_pct, "matched": ats.matched,
                    "missing": ats.missing, "ats": f"{ats.match_pct*100:.0f}% — {ats.flag}",
                    "cv_bytes": open(res.docx_path, "rb").read(),
                    "cv_pdf": open(res.pdf_path, "rb").read() if res.pdf_path else None,
                    "pages": res.page_count,
                })
                st.session_state["tlr"] = r
            st.rerun()

        with st.expander("Original → tailored (what changed)"):
            for bf, af in zip(r["before"], r["after"]):
                st.markdown(
                    f"<div class='card'><div class='role'>❌ {bf}</div>"
                    f"<div class='firm' style='margin-top:.45rem'>✅ {af}</div></div>",
                    unsafe_allow_html=True,
                )
        with st.expander("Cover letter"):
            st.text_area("cover", r["cover"], height=260, label_visibility="collapsed")

# --- admin panel: who signed up + free upgrade toggle ------------------------
if tab_admin is not None:
    with tab_admin:
        st.subheader("🛠️ Admin — signups & access")
        people = acc.list_users()
        total_u = len(people)
        pros = sum(1 for u in people if u["tier"] in ("pro", "admin"))
        m1, m2, m3 = st.columns(3)
        m1.markdown(f"<div class='metric'><div class='v'>{total_u}</div><div class='l'>Sign-ups</div></div>", unsafe_allow_html=True)
        m2.markdown(f"<div class='metric'><div class='v'>{pros}</div><div class='l'>Pro / admin</div></div>", unsafe_allow_html=True)
        m3.markdown(f"<div class='metric'><div class='v'>{total_u - pros}</div><div class='l'>Free</div></div>", unsafe_allow_html=True)
        st.write("")
        st.caption("Flip anyone to Pro (free) or back. Changes apply the moment they reload.")
        for person in people:
            c1, c2, c3 = st.columns([3, 1.4, 1.4])
            c1.markdown(f"**{person['email']}**  \n<span class='role'>joined {person['created_at'][:10]}</span>",
                        unsafe_allow_html=True)
            c2.markdown(f"<span class='pill' style='background:"
                        f"{'#a855f7' if person['tier']=='admin' else '#10b981' if person['tier']=='pro' else '#64748b'};"
                        f"margin-top:.4rem;display:inline-block'>{person['tier'].upper()}</span>",
                        unsafe_allow_html=True)
            if person["tier"] == "free":
                if c3.button("⬆ Make Pro (free)", key=f"up_{person['email']}"):
                    acc.set_tier(person["email"], "pro")
                    st.rerun()
            elif person["tier"] == "pro":
                if c3.button("⬇ Back to Free", key=f"down_{person['email']}"):
                    acc.set_tier(person["email"], "free")
                    st.rerun()
            else:
                c3.caption("admin")
